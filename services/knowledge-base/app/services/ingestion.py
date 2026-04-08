from __future__ import annotations

import mimetypes
import os
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from uuid import uuid4

from docx import Document as WordDocument
from pypdf import PdfReader
from sqlalchemy import delete, select
from sqlalchemy.orm import Session

from app.core.config import settings
from app.db.models import KnowledgeChunk, NormativeSource
from app.services.audit import record_audit_event
from app.services.embeddings import embed_text


@dataclass
class ParsedDocument:
    text: str
    metadata: dict[str, Any]


def _clean_text(value: str | None) -> str:
    if not value:
        return ""
    return re.sub(r"\s+", " ", value).strip()


def ensure_storage_dir() -> Path:
    path = Path(settings.KNOWLEDGE_BASE_STORAGE_DIR)
    path.mkdir(parents=True, exist_ok=True)
    return path


def _guess_mime(file_name: str, explicit: str | None = None) -> str:
    if explicit:
        return explicit
    return mimetypes.guess_type(file_name)[0] or "application/octet-stream"


def save_upload(file_name: str, content: bytes) -> tuple[str, int]:
    storage_dir = ensure_storage_dir()
    target_name = f"{uuid4().hex}_{Path(file_name).name}"
    target_path = storage_dir / target_name
    target_path.write_bytes(content)
    return str(target_path), len(content)


def _parse_pdf(local_path: str) -> ParsedDocument:
    reader = PdfReader(local_path)
    fragments: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = _clean_text(page.extract_text())
        if not text:
            continue
        fragments.append(f"[PAGE {index}] {text}")
    return ParsedDocument(
        text="\n\n".join(fragments).strip(),
        metadata={"page_count": len(reader.pages), "parser": "pypdf"},
    )


def _parse_docx(local_path: str) -> ParsedDocument:
    document = WordDocument(local_path)
    lines: list[str] = []
    paragraph_count = 0
    table_count = 0

    for paragraph in document.paragraphs:
        text = _clean_text(paragraph.text)
        if not text:
            continue
        paragraph_count += 1
        lines.append(f"[P{paragraph_count}] {text}")

    for table in document.tables:
        table_count += 1
        lines.append(f"[T{table_count}] TABLE START")
        for row_index, row in enumerate(table.rows, start=1):
            cells = [_clean_text(cell.text) or "—" for cell in row.cells]
            if not any(cell != "—" for cell in cells):
                continue
            lines.append(f"[T{table_count}R{row_index}] {' | '.join(cells)}")
        lines.append(f"[T{table_count}] TABLE END")

    return ParsedDocument(
        text="\n".join(lines).strip(),
        metadata={"paragraph_count": paragraph_count, "table_count": table_count, "parser": "python-docx"},
    )


def _parse_text(local_path: str) -> ParsedDocument:
    content = Path(local_path).read_text(encoding="utf-8", errors="ignore")
    return ParsedDocument(text=content.strip(), metadata={"parser": "plain-text"})


def parse_document(local_path: str, file_name: str, mime_type: str | None = None) -> ParsedDocument:
    suffix = Path(file_name).suffix.lower()
    normalized_mime = _guess_mime(file_name, mime_type)
    if suffix == ".pdf" or normalized_mime == "application/pdf":
        return _parse_pdf(local_path)
    if suffix in {".docx", ".doc"} or normalized_mime.endswith("wordprocessingml.document"):
        return _parse_docx(local_path)
    return _parse_text(local_path)


def _page_markers(text: str) -> tuple[int | None, int | None]:
    pages = [int(item) for item in re.findall(r"\[PAGE\s+(\d+)\]", text)]
    if not pages:
        return None, None
    return min(pages), max(pages)


def chunk_text(text: str) -> list[dict[str, Any]]:
    normalized = text.strip()
    if not normalized:
        return []
    chunk_size = settings.KNOWLEDGE_BASE_CHUNK_SIZE_CHARS
    overlap = settings.KNOWLEDGE_BASE_CHUNK_OVERLAP_CHARS
    chunks: list[dict[str, Any]] = []
    start = 0
    index = 0
    length = len(normalized)
    while start < length:
        end = min(length, start + chunk_size)
        if end < length:
            split_at = normalized.rfind("\n\n", start, end)
            if split_at > start + 300:
                end = split_at
        raw_chunk = normalized[start:end].strip()
        if raw_chunk:
            page_start, page_end = _page_markers(raw_chunk)
            chunks.append(
                {
                    "chunk_index": index,
                    "text": raw_chunk,
                    "token_count": max(1, len(raw_chunk.split())),
                    "page_start": page_start,
                    "page_end": page_end,
                    "metadata_json": {"length": len(raw_chunk)},
                }
            )
            index += 1
        if end >= length:
            break
        start = max(end - overlap, start + 1)
    return chunks


def _next_source_version(db: Session, source_key: str) -> int:
    versions = db.scalars(
        select(NormativeSource.version).where(NormativeSource.source_key == source_key)
    ).all()
    return (max(versions) if versions else 0) + 1


def _next_source_version_for_project(db: Session, project_key: str, source_key: str) -> int:
    versions = db.scalars(
        select(NormativeSource.version).where(
            NormativeSource.project_key == project_key,
            NormativeSource.source_key == source_key,
        )
    ).all()
    return (max(versions) if versions else 0) + 1


def ingest_normative_document(
    db: Session,
    *,
    project_key: str,
    source_key: str,
    title: str,
    source_type: str,
    jurisdiction: str | None,
    status: str,
    effective_from: Any,
    effective_to: Any,
    summary: str | None,
    metadata_json: dict[str, Any] | None,
    file_name: str,
    file_bytes: bytes,
    actor: str = "admin",
) -> NormativeSource:
    storage_path, file_size = save_upload(file_name, file_bytes)
    parsed = parse_document(storage_path, file_name)
    version = _next_source_version_for_project(db, project_key, source_key)
    source = NormativeSource(
        project_key=project_key,
        source_key=source_key,
        version=version,
        title=title,
        source_type=source_type,
        jurisdiction=jurisdiction,
        status=status,
        is_published=status == "published",
        effective_from=effective_from,
        effective_to=effective_to,
        file_name=file_name,
        mime_type=_guess_mime(file_name),
        file_size=file_size,
        storage_path=storage_path,
        content_text=parsed.text,
        summary=summary,
        metadata_json={**(metadata_json or {}), **parsed.metadata},
    )
    db.add(source)
    db.flush()

    chunks = chunk_text(parsed.text)
    for chunk in chunks:
        db.add(
            KnowledgeChunk(
                project_key=source.project_key,
                source_id=source.id,
                source_key=source.source_key,
                source_type=source.source_type,
                chunk_index=chunk["chunk_index"],
                page_start=chunk["page_start"],
                page_end=chunk["page_end"],
                text=chunk["text"],
                token_count=chunk["token_count"],
                embedding=embed_text(chunk["text"]),
                metadata_json=chunk["metadata_json"],
                is_published=source.is_published,
            )
        )

        record_audit_event(
        db,
        project_key=source.project_key,
        entity_type="normative_source",
        entity_id=source.id,
        action="ingest",
        actor=actor,
        snapshot={
            "source_key": source.source_key,
            "version": source.version,
            "title": source.title,
            "status": source.status,
            "chunk_count": len(chunks),
        },
    )
    return source


def reindex_source_chunks(db: Session, source: NormativeSource, actor: str = "system") -> int:
    if not source.content_text:
        return 0
    db.execute(delete(KnowledgeChunk).where(KnowledgeChunk.source_id == source.id))
    chunks = chunk_text(source.content_text)
    for chunk in chunks:
        db.add(
            KnowledgeChunk(
                project_key=source.project_key,
                source_id=source.id,
                source_key=source.source_key,
                source_type=source.source_type,
                chunk_index=chunk["chunk_index"],
                page_start=chunk["page_start"],
                page_end=chunk["page_end"],
                text=chunk["text"],
                token_count=chunk["token_count"],
                embedding=embed_text(chunk["text"]),
                metadata_json=chunk["metadata_json"],
                is_published=source.is_published,
            )
        )
    record_audit_event(
        db,
        project_key=source.project_key,
        entity_type="normative_source",
        entity_id=source.id,
        action="reindex",
        actor=actor,
        snapshot={"source_key": source.source_key, "version": source.version, "chunk_count": len(chunks)},
    )
    return len(chunks)
